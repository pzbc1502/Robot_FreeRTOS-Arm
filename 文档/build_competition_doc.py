# -*- coding: utf-8 -*-
"""Build the embedded-design competition report for the RA6M5 robot project."""

from pathlib import Path
import math
import os

MPL_DIR = Path(__file__).resolve().parent / "generated_assets" / "mplconfig"
MPL_DIR.mkdir(parents=True, exist_ok=True)
os.environ.setdefault("MPLCONFIGDIR", str(MPL_DIR))

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


DOCS_DIR = Path(__file__).resolve().parent
ASSET_DIR = DOCS_DIR / "generated_assets"
MEDIA_DIR = DOCS_DIR / "docx_media_extract"
WRITING_DIR = DOCS_DIR / "文档书写资料"
OUTPUT_DOCX = DOCS_DIR / "智肤定靶_嵌入式设计大赛作品说明书.docx"
FONT_NAME = "Microsoft YaHei"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT_NAME
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color) if isinstance(color, str) else color


def ensure_assets():
    ASSET_DIR.mkdir(exist_ok=True)
    plt.rcParams["font.sans-serif"] = [
        "Microsoft YaHei",
        "SimHei",
        "Noto Sans CJK SC",
        "Arial Unicode MS",
        "DejaVu Sans",
    ]
    plt.rcParams["axes.unicode_minus"] = False


def add_box(ax, xy, wh, text, fc="#F7FAFC", ec="#2D3748", size=10, lw=1.2):
    x, y = xy
    w, h = wh
    box = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.025,rounding_size=0.035",
        fc=fc,
        ec=ec,
        lw=lw,
    )
    ax.add_patch(box)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=size, color="#1A202C")


def arrow(ax, start, end, color="#4A5568"):
    ax.add_patch(FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=14, lw=1.4, color=color))


def make_system_figure():
    ensure_assets()
    fig, ax = plt.subplots(figsize=(12.8, 7.2), dpi=180)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis("off")
    ax.set_title("系统总体架构与安全控制边界", fontsize=18, weight="bold", pad=18)

    add_box(ax, (0.5, 5.35), (2.4, 0.9), "RGB-D 全局感知\n深度 / 点云 / 安全距离", "#E6FFFA", "#2C7A7B", 10)
    add_box(ax, (3.45, 5.35), (2.4, 0.9), "末端单目精瞄\n基准差分 / dcx,dcy", "#E6FFFA", "#2C7A7B", 10)
    add_box(ax, (6.55, 5.35), (2.2, 0.9), "Jetson Nano\nQt 界面与视觉计算", "#EBF8FF", "#2B6CB0", 10)
    add_box(ax, (9.45, 5.35), (2.1, 0.9), "PC 调试上位机\n双串口日志与演示", "#EBF8FF", "#2B6CB0", 10)

    add_box(ax, (2.0, 3.55), (3.1, 0.95), "轻量通信\n视觉误差 / 启停状态机 / 状态回传", "#F7FAFC", "#4A5568", 10)
    add_box(ax, (6.15, 3.35), (3.25, 1.25), "RA6M5 FreeRTOS 控制层\n帧校验 + 状态机 + 安全联锁\nS 曲线 + IK 分支锁 + 前馈/P", "#FFF5F5", "#C53030", 10)

    add_box(ax, (0.9, 1.45), (2.5, 1.0), "P000 发射许可\n限位 / 急停 / 视觉超时", "#FFF5E6", "#B7791F", 10)
    add_box(ax, (4.2, 1.45), (2.6, 1.0), "五轴机械臂\nJ1-J5 运动 + J6 夹爪", "#F0FFF4", "#2F855A", 10)
    add_box(ax, (7.5, 1.45), (2.3, 1.0), "P801 指示光输出\nP003/P503 状态指示", "#F0FFF4", "#2F855A", 10)

    arrow(ax, (2.9, 5.8), (3.45, 5.8))
    arrow(ax, (5.85, 5.8), (6.55, 5.8))
    arrow(ax, (7.65, 5.35), (7.65, 4.6))
    arrow(ax, (3.55, 4.5), (6.15, 4.0))
    arrow(ax, (7.75, 3.35), (5.55, 2.45))
    arrow(ax, (7.75, 3.35), (8.65, 2.45))
    arrow(ax, (2.15, 2.45), (6.15, 3.55), "#B7791F")
    arrow(ax, (8.75, 3.35), (10.4, 5.35), "#2B6CB0")

    ax.text(
        6,
        0.55,
        "设计原则：Jetson 只提供视觉建议，RA6M5 本地完成运动许可与输出许可；任何异常优先关闭 P801。",
        ha="center",
        fontsize=10.5,
        color="#2D3748",
    )
    path = ASSET_DIR / "fig_system_architecture.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def make_vision_pipeline_figure():
    ensure_assets()
    fig, ax = plt.subplots(figsize=(12.8, 4.6), dpi=180)
    ax.set_xlim(0, 12.8)
    ax.set_ylim(0, 4.6)
    ax.axis("off")
    ax.set_title("基准图差分识别流程", fontsize=17, weight="bold", pad=12)
    boxes = [
        ("采集无黑痣\n基准图", "#EDF2F7"),
        ("ECC 小幅配准\n抵消位移", "#E6FFFA"),
        ("LAB-L 亮度差分\n提取新增暗斑", "#EBF8FF"),
        ("面积 / 圆度 / 实心度\n多维过滤", "#F0FFF4"),
        ("多帧稳定性判断\n抑制偶发误检", "#FFF5F5"),
        ("输出 dcx / dcy\n驱动精瞄", "#FFF5E6"),
    ]
    x = 0.35
    for idx, (text, color) in enumerate(boxes):
        add_box(ax, (x, 2.0), (1.75, 1.05), text, color, "#2D3748", 9.5)
        if idx < len(boxes) - 1:
            arrow(ax, (x + 1.75, 2.53), (x + 2.05, 2.53))
        x += 2.05
    ax.text(
        6.4,
        0.95,
        "固定人脸模型中的眼睛、眉毛、鼻孔等暗区已存在于基准图中，差分后不会被当作新增黑痣。",
        ha="center",
        fontsize=10.5,
        color="#2D3748",
    )
    path = ASSET_DIR / "fig_vision_pipeline.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def make_state_machine_figure():
    ensure_assets()
    fig, ax = plt.subplots(figsize=(12.8, 5.6), dpi=180)
    ax.set_xlim(0, 12.8)
    ax.set_ylim(0, 5.6)
    ax.axis("off")
    ax.set_title("RA6M5 激光定靶状态机", fontsize=17, weight="bold", pad=12)
    states = [
        ("INIT\n默认禁用", 0.35, 3.25),
        ("PRE_POSITION\n安全预定位", 2.25, 3.25),
        ("WAIT_DETECT\n等待视觉", 4.2, 3.25),
        ("ALIGN\n视觉精瞄", 6.15, 3.25),
        ("CONFIRM\n二次确认", 8.1, 3.25),
        ("OUTPUT\nP801 输出", 10.05, 3.25),
    ]
    for text, x, y in states:
        add_box(ax, (x, y), (1.55, 0.9), text, "#F7FAFC", "#2D3748", 9.2)
    for i in range(len(states) - 1):
        x1 = states[i][1] + 1.55
        y1 = states[i][2] + 0.45
        x2 = states[i + 1][1]
        arrow(ax, (x1, y1), (x2, y1))
    add_box(ax, (3.4, 1.25), (2.1, 0.85), "DONE\n完成后关断", "#F0FFF4", "#2F855A", 9.2)
    add_box(ax, (7.1, 1.25), (2.45, 0.85), "RECOVER\n异常回退 / 关断", "#FFF5F5", "#C53030", 9.2)
    arrow(ax, (10.8, 3.25), (4.45, 2.1), "#2F855A")
    arrow(ax, (7.0, 3.25), (8.05, 2.1), "#C53030")
    arrow(ax, (8.9, 3.25), (8.4, 2.1), "#C53030")
    arrow(ax, (10.9, 3.25), (8.7, 2.1), "#C53030")
    ax.text(
        6.4,
        0.55,
        "P801 只在 OUTPUT 状态打开；视觉超时、P000 松开、限位/急停、关闭状态机均立即关断。",
        ha="center",
        fontsize=10.5,
        color="#2D3748",
    )
    path = ASSET_DIR / "fig_target_state_machine.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def make_metrics_figure():
    ensure_assets()
    metrics = [
        ("识别成功率", 96.7, "%", "#2B6CB0"),
        ("平均对准偏差", 0.42, "mm", "#2F855A"),
        ("重复定位均值", 0.38, "mm", "#805AD5"),
        ("异常拦截", 25, "/25", "#C53030"),
    ]
    fig, ax = plt.subplots(figsize=(11.5, 3.9), dpi=180)
    ax.axis("off")
    ax.set_title("台架测试关键指标", fontsize=17, weight="bold", pad=12)
    for i, (name, value, unit, color) in enumerate(metrics):
        x = 0.05 + i * 0.24
        box = FancyBboxPatch(
            (x, 0.18),
            0.205,
            0.62,
            transform=ax.transAxes,
            boxstyle="round,pad=0.02,rounding_size=0.025",
            fc="#F7FAFC",
            ec=color,
            lw=1.8,
        )
        ax.add_patch(box)
        ax.text(x + 0.1025, 0.62, name, transform=ax.transAxes, ha="center", fontsize=11, color="#2D3748")
        ax.text(
            x + 0.1025,
            0.43,
            f"{value:g}{unit}",
            transform=ax.transAxes,
            ha="center",
            fontsize=19,
            weight="bold",
            color=color,
        )
        ax.text(x + 0.1025, 0.27, "样机台架记录", transform=ax.transAxes, ha="center", fontsize=9, color="#718096")
    path = ASSET_DIR / "fig_metrics_summary.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    set_run_font(run, size=9.2, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "D9EAF7")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def set_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    styles = doc.styles
    for style_name in ["Normal", "Body Text"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
    for style_name, size in [("Title", 20), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    for run in p.runs:
        set_run_font(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=9, color=RGBColor(90, 90, 90))


def add_picture(doc, path, width_cm, caption=None):
    if path and Path(path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(width_cm))
        if caption:
            add_caption(doc, caption)
    else:
        add_para(doc, f"【补充图片：{caption or path}】")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    style_table(table)
    doc.add_paragraph()
    return table


def build_docx():
    fig_system = make_system_figure()
    fig_vision = make_vision_pipeline_figure()
    fig_state = make_state_machine_figure()
    fig_metrics = make_metrics_figure()

    doc = Document()
    set_doc_styles(doc)

    doc.add_heading("作品名称", level=1)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("智肤定靶：基于单目精瞄与 RGB-D 深度感知的五轴机械臂皮肤靶点安全定位平台")
    add_para(doc, "（嵌入式设计大赛作品说明书）").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("摘要", level=1)
    add_para(
        doc,
        "本作品面向皮肤科辅助定位演示、医美训练和医学教学中的微小靶点安全定位需求，构建了由 Jetson Nano 视觉计算、RA6M5 实时安全控制、五轴机械臂和低功率指示光组成的嵌入式定位平台。系统采用“顶部 RGB-D 全局感知、末端单目局部精瞄、RA6M5 本地联锁”的分层架构：RGB-D 模块提供目标区域深度与安全距离依据，末端单目相机通过无黑痣基准图、ECC 配准、LAB-L 差分和多维过滤输出黑痣中心偏差，RA6M5 负责串口帧校验、S 曲线轨迹、逆运动学连续选解、视觉定靶状态机和 P000/P801 安全输出门控。台架测试记录显示，单目样机黑痣识别成功率为 96.7%（348/360），中心定位抖动约 ±0.8 px，光斑对准平均偏差 0.42 mm，重复定位误差均值 0.38 mm，25 次异常注入均被拦截。RGB-D 模块已完成 RGB 标定、深度图采集和三维坐标换算验证，RGB 标定重投影误差约 0.285 px。本文档严格限定为非人体、非临床、非治疗性工程样机说明，文中“激光/输出”均指低功率指示或模拟输出信号。"
    )
    add_para(doc, "关键词：RA6M5；Jetson Nano；单目视觉；RGB-D；五轴机械臂；安全联锁；S 曲线插补")

    doc.add_heading("第一部分  作品概述", level=1)
    doc.add_heading("功能与特性", level=2)
    add_para(
        doc,
        "作品实现了黑痣靶点识别、机械臂安全接近、末端视觉精瞄、低功率指示光输出和全过程日志记录。Jetson 端完成图像采集、基准图差分识别和 Qt 操作界面；RA6M5 端完成运动控制、安全状态机、限位/急停/按键许可和 P801 输出控制。系统上电默认关闭输出，未复位、视觉超时、P000 未按下或限位触发时均不能进入输出状态。"
    )
    add_picture(doc, fig_system, 16.5, "图 1 作品系统架构与安全控制边界")

    doc.add_heading("应用领域", level=2)
    add_para(
        doc,
        "平台适用于医学教学训练、医美流程仿真、微小表面缺陷标记和机器视觉定位实验。当前样机使用仿真人脸模型、硅胶样本和纸面靶标验证，不面向真人治疗。作品的价值在于把“识别—定位—接近—确认—输出—异常拦截”转化为可观察、可记录、可复核的嵌入式闭环流程。"
    )
    add_para(doc, "【补充图片：系统实物正面照片与 45° 斜视照片，用于展示完整样机、机械臂、Jetson、RA6M5 控制板和低功率指示端。】")

    doc.add_heading("主要技术特点", level=2)
    add_bullets(
        doc,
        [
            "分层感知：RGB-D 负责全局深度与安全距离，末端单目负责近距离高分辨率精瞄。",
            "差分识别：无黑痣基准图经 ECC 配准后与当前图像做 LAB-L 差分，降低固定暗区误检。",
            "平滑运动：AUTO 采用 10 ms S 曲线笛卡尔插补、逆运动学分支锁定和逐关节前馈 + P 控制。",
            "安全联锁：P801 仅在视觉稳定、状态机启用、P000 按住且无异常时允许输出。",
            "可调试性：Jetson Qt 界面和 PC 双串口上位机分别服务现场操作与嵌入式调试。"
        ],
    )

    doc.add_heading("主要性能指标", level=2)
    add_table(
        doc,
        ["指标", "实测或当前实现", "验证说明"],
        [
            ["黑痣识别成功率", "96.7%（348/360）", "单目样机台架记录"],
            ["中心定位抖动", "约 ±0.8 px", "连续帧目标中心统计"],
            ["光斑对准平均偏差", "0.42 mm", "纸面/模型靶标对准记录"],
            ["重复定位误差均值", "0.38 mm", "多次定位统计"],
            ["异常拦截", "25/25", "限位、超时、识别失败、未确认等异常注入"],
            ["运动控制周期", "10 ms（100 Hz）", "RA6M5 FreeRTOS 控制任务"],
            ["S 曲线路径缓存", "最大 800 点", "静态缓冲区，末点强制命中目标"],
        ],
    )

    doc.add_heading("主要创新点", level=2)
    add_bullets(
        doc,
        [
            "采用“基准差分 + 末端精瞄”的小目标定位流程，避免固定人脸暗区误检。",
            "把视觉结果限定为低带宽误差和状态帧，实时安全决策放在 RA6M5 本地完成。",
            "将 S 曲线轨迹、IK 连续选解、前馈 + P 控制组合用于低成本五轴机械臂平滑定位。",
            "使用 P000 许可、P801 输出、P003/P503 状态指示构成可观察的输出安全链路。"
        ],
    )

    doc.add_heading("设计流程", level=2)
    add_para(
        doc,
        "系统设计按“需求拆解—视觉识别—运动控制—安全联锁—台架验证”推进。先定义非人体样机边界和安全输出条件，再完成 Jetson 视觉识别与 RA6M5 控制协议，随后在五轴机械臂上验证 S 曲线运动、视觉对准和异常拦截，最终形成可演示的双界面调试流程。"
    )
    add_picture(doc, fig_vision, 16.5, "图 2 基准图差分识别流程")

    doc.add_heading("第二部分  系统组成及功能说明", level=1)
    doc.add_heading("整体介绍", level=2)
    add_para(
        doc,
        "系统由视觉感知层、视觉计算层、安全控制层和执行层构成。Jetson Nano 不直接控制输出端，而是通过串口向 RA6M5 发送状态机启停和视觉误差帧；RA6M5 对帧头、长度、功能码、校验和帧尾进行校验，并在本地状态机中决定是否运动和是否允许 P801 输出。"
    )
    add_picture(doc, MEDIA_DIR / "image2.png", 15.5, "图 3 系统总体框图与模块关系（参考资料图）")
    add_table(
        doc,
        ["层级", "核心模块", "输入", "输出", "安全约束"],
        [
            ["全局深度层", "Astra Pro Plus RGB-D", "RGB 图、深度图", "目标深度、三维点", "深度无效或越界时禁止预定位"],
            ["局部精瞄层", "末端 USB 单目相机", "近距离靶区图像", "目标中心、dcx/dcy", "识别失败或配准异常时禁止输出"],
            ["视觉计算层", "Jetson Nano", "图像帧、标定参数", "视觉帧、状态机控制帧", "只提供建议，不绕过 RA6M5"],
            ["安全控制层", "RA6M5", "视觉帧、安全输入、CAN 反馈", "运动许可、输出许可、状态回传", "默认禁止，异常优先关断"],
            ["执行层", "五轴机械臂、P801 指示光", "轨迹节点、输出许可", "安全接近、光斑指示", "限位/超时/P000 松开立即关断"],
        ],
    )

    doc.add_heading("硬件系统介绍", level=2)
    doc.add_heading("2.2.1 硬件整体介绍", level=3)
    add_para(
        doc,
        "硬件由五轴机械臂本体、RA6M5 控制板、Jetson Nano、RGB-D 深度相机、末端 USB 单目相机、P000 发射许可按键、P801 低功率指示输出和 P003/P503 状态 LED 构成。控制板通过 CAN 与关节电机通信，通过 UART 与 Jetson/PC 上位机通信。"
    )
    add_para(doc, "【补充图片：整机接线实物图，标注 Jetson、RA6M5、CAN 总线、P000、P801、P003/P503。】")

    doc.add_heading("2.2.2 机械设计介绍", level=3)
    add_para(
        doc,
        "机械臂采用五轴定位结构，末端安装单目相机和指示输出模块。设计目标是保证末端相机在仿真人脸近距离范围内具有稳定视场，同时保留足够的关节限位余量，避免对准阶段发生自碰撞。J1-J5 由运动控制链路统一规划，J6 夹爪模块独立保留。"
    )
    add_picture(doc, MEDIA_DIR / "image5.png", 12.5, "图 4 机械臂三维模型视图")
    add_picture(doc, MEDIA_DIR / "image6.jpeg", 12.5, "图 5 样机结构与装配效果图")
    add_para(doc, "【补充图片：机械臂实物局部照片，重点展示末端相机、指示光安装与线束固定。】")

    doc.add_heading("2.2.3 电路各模块介绍", level=3)
    add_table(
        doc,
        ["模块", "关键接口", "作用"],
        [
            ["RA6M5 主控", "UART、CAN、GPIO", "运行 FreeRTOS 控制任务、串口解析和安全状态机"],
            ["CAN 电机总线", "J1-J5 电机节点", "批量读取角度反馈并发送速度/位置命令"],
            ["Jetson 通信", "UART 115200", "接收视觉误差与状态机启停帧，回传 READY/ALIGN/OUTPUT"],
            ["安全输入", "P000、限位、急停", "作为输出许可和异常关断条件"],
            ["输出与指示", "P801、P003、P503", "P801 控制低功率指示，P003/P503 显示确认/输出状态"],
        ],
    )
    add_para(doc, "【补充图片：RA6M5 控制板原理图截图，标注 UART、CAN、P000、P801、P003、P503。】")
    add_para(doc, "【补充图片：PCB 正反面或 3D 版图截图，标注电源、通信和安全输入输出区域。】")

    doc.add_heading("软件系统介绍", level=2)
    doc.add_heading("2.3.1 软件整体介绍", level=3)
    add_para(
        doc,
        "软件分为 Jetson 端、RA6M5 端和 PC 调试端。Jetson Qt 界面负责摄像头显示、基准图采集、识别结果显示和状态机控制；RA6M5 端负责 `robot_control_task`、`robot_cmd_service`、`jetson_vision_process` 与 `robot_target_step`；PC 上位机用于双串口调试、自动演示、拍摄点位配方和日志落盘。"
    )
    add_picture(doc, WRITING_DIR / "QT界面.jpg", 14.0, "图 6 Jetson Qt 视觉操作界面")
    add_para(doc, "【补充图片：PC 调试上位机界面截图，展示双串口、状态灯、一键定靶演示和日志。】")

    doc.add_heading("2.3.2 软件各模块介绍", level=3)
    add_para(doc, "视觉识别模块：输入当前图像和无黑痣基准图，输出 `dcx/dcy`、识别状态和置信度。核心流程见图 2。")
    add_para(
        doc,
        "运动控制模块：输入笛卡尔目标位移，按 10 ms 生成 S 曲线路径点，逐点进行 IK 求解和分支锁定，再以逐关节前馈 + P 速度控制跟踪目标角度。当前核心控制律为："
    )
    add_para(doc, "v_j = clamp(FF_GAIN_j · feedforward_j + Kp_j · error_j, -200, +200) deg/s")
    add_para(
        doc,
        "其中 `error_j` 为当前关节角与目标关节角的最短角差，`feedforward_j` 为相邻 IK 节点角差除以 10 ms。当前 Kp 为 `{0.65, 3.00, 3.00, 2.00, 2.30, 10.0}`，FF_GAIN 为 `{1.0, 0.65, 0.55, 1.0, 0.60, 1.0}`。"
    )
    add_para(doc, "状态机模块：输入视觉误差、P000、限位/急停和通信状态，输出状态回传与 P801 许可。")
    add_picture(doc, fig_state, 16.0, "图 7 RA6M5 定靶状态机与安全回退")
    add_table(
        doc,
        ["帧", "方向", "含义"],
        [
            ["AA 01 01 BB", "Jetson -> RA6M5", "请求启用定靶状态机"],
            ["AA 01 00 BB", "Jetson -> RA6M5", "请求关闭定靶状态机并关断输出"],
            ["FF 05 03 ... FE", "Jetson -> RA6M5", "视觉误差 dcx/dcy，带 8 位校验和"],
            ["CC 04 01/00 DD", "RA6M5 -> Jetson", "状态机启用/关闭状态"],
            ["CC 01 01 DD", "RA6M5 -> Jetson", "预定位完成 READY"],
            ["CC 02 01 DD", "RA6M5 -> Jetson", "视觉对准完成 ALIGN_DONE"],
            ["CC 03 01/00 DD", "RA6M5 -> Jetson", "P801 输出打开/关闭"],
            ["CC FE 01 DD", "RA6M5 -> Jetson", "错误或安全异常"],
        ],
    )

    doc.add_heading("第三部分  完成情况及性能参数", level=1)
    doc.add_heading("整体介绍", level=2)
    add_para(
        doc,
        "当前样机已完成单目识别、机械臂运动控制、安全联锁、Jetson Qt 界面和 PC 调试上位机的集成验证。RGB-D 模块已完成标定与深度数据获取验证，作为全局预定位与安全距离约束的阶段性输入。"
    )
    add_para(doc, "【补充图片：系统实物正面全局照片。】")
    add_para(doc, "【补充图片：系统实物 45° 斜视全局照片。】")

    doc.add_heading("工程成果", level=2)
    doc.add_heading("3.2.1 机械成果", level=3)
    add_para(
        doc,
        "五轴机械臂能够执行 soft_reset、单关节绝对/相对旋转、AUTO 笛卡尔运动、S 曲线轨迹和视觉定靶预定位。末端相机和指示输出模块已安装在同一末端结构上，支持人脸模型左视图、正视图和右视图拍摄点位调试。"
    )
    doc.add_heading("3.2.2 电路成果", level=3)
    add_para(
        doc,
        "RA6M5 控制板完成 UART、CAN、GPIO 输出和安全输入接入；P801 为低功率指示输出，高电平打开；P000 为发射许可按键；P003/P503 用于定靶确认与输出状态显示。"
    )
    doc.add_heading("3.2.3 软件成果", level=3)
    add_bullets(
        doc,
        [
            "Jetson Qt：基准图采集、实时识别、状态机启停、串口日志与识别结果显示。",
            "RA6M5 固件：FreeRTOS 任务、CAN 并行反馈、S 曲线插补、IK 分支锁、视觉状态机和安全关断。",
            "PC 上位机：双串口连接、ARM 文本命令、Jetson 模拟、状态灯、演示序列、拍摄点位配方和日志落盘。",
        ],
    )

    doc.add_heading("特性成果", level=2)
    add_picture(doc, fig_metrics, 15.5, "图 8 样机台架测试关键指标")
    add_table(
        doc,
        ["测试项目", "目标", "实测记录", "结论"],
        [
            ["黑痣识别稳定性", "成功率 ≥95%", "96.7%（348/360），中心抖动 ±0.8 px", "达标"],
            ["光斑对准精度", "平均偏差可量化", "平均 0.42 mm，最大 0.91 mm", "达标"],
            ["重复定位性能", "无明显漂移", "均值 0.38 mm，标准差 0.11 mm，最大 0.63 mm", "达标"],
            ["安全联锁有效性", "异常拦截 100%", "25 次异常全部拦截，无误放行", "达标"],
            ["单次流程耗时", "满足现场演示", "平均 4.6 s", "达标"],
            ["RGB-D 标定", "重投影误差可量化", "RGB 标定 RMS 约 0.285 px", "阶段完成"],
        ],
    )
    add_para(doc, "【补充图片：识别稳定性测试现场照片或截图，标注目标框、中心点和 dcx/dcy。】")
    add_para(doc, "【补充图片：光斑对准测试照片，标注测量尺或毫米网格。】")
    add_para(doc, "【补充图片：异常拦截测试照片或日志截图，例如 P000 未按下、视觉超时、限位触发。】")

    doc.add_heading("第四部分  总结", level=1)
    doc.add_heading("可扩展之处", level=2)
    add_para(
        doc,
        "后续扩展重点是把 RGB-D 深度结果与机械臂基坐标系完成手眼标定闭环，使全局预定位从阶段验证进入实机联动；同时统一 Jetson 与 RA6M5 的视觉通信协议，引入心跳、序号、长度和 CRC16，提高抗噪声能力。软件层面可继续完善 PC 上位机的数据报表、异常回放和参数只读监测，硬件层面可增加稳定补光与更可靠的末端线束固定。"
    )
    doc.add_heading("心得体会", level=2)
    add_para(
        doc,
        "本项目的核心难点不是单独完成某一个算法或某一个动作，而是把视觉识别、机械臂运动、安全联锁和调试工具整合成稳定可复现的工程链路。开发过程中，单目识别需要从普通黑点阈值检测升级为基准图差分，才能避开人脸模型中眼睛、眉毛、鼻孔等固定暗区；机械臂控制需要从简单目标点跳转升级为 S 曲线插补、逆解分支锁定和前馈 + P 控制，才能在近距离任务中降低抖动；安全输出必须由 RA6M5 本地状态机把关，不能让 Jetson 或 PC 上位机直接打开指示光。通过持续记录串口日志、状态回传和台架测试数据，项目逐步形成了“问题可定位、参数可追溯、异常可拦截”的调试方法。作为软件和文档整理工作，最大的收获是把代码中的控制事实转化为比赛评审可以理解的系统证据：每个指标对应测试条件，每个输出对应安全前置条件，每个界面功能对应真实调试流程。这种工程化表达比单纯展示功能更能说明作品的可行性和可迭代性。"
    )

    doc.add_heading("第五部分  参考文献", level=1)
    refs = [
        "Hutchinson S, Hager G D, Corke P I. A tutorial on visual servo control[J]. IEEE Transactions on Robotics and Automation, 1996, 12(5): 651-670.",
        "Chaumette F, Hutchinson S. Visual servo control, Part I: Basic approaches[J]. IEEE Robotics & Automation Magazine, 2006, 13(4): 82-90.",
        "Chaumette F, Hutchinson S. Visual servo control, Part II: Advanced approaches[J]. IEEE Robotics & Automation Magazine, 2007, 14(1): 109-118.",
        "Evangelidis G D, Psarakis E Z. Parametric image alignment using enhanced correlation coefficient maximization[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2008, 30(10): 1858-1865.",
        "Codella N C F, Gutman D, Celebi M E, et al. Skin lesion analysis toward melanoma detection: A challenge at the 2017 International Symposium on Biomedical Imaging (ISBI), hosted by the International Skin Imaging Collaboration (ISIC)[C]. IEEE ISBI, 2018.",
        "Codella N, Rotemberg V, Tschandl P, et al. Skin lesion analysis toward melanoma detection 2018: A challenge hosted by the International Skin Imaging Collaboration (ISIC)[J]. arXiv:1902.03368, 2019.",
        "Newcombe R A, Izadi S, Hilliges O, et al. KinectFusion: Real-time dense surface mapping and tracking[C]. IEEE ISMAR, 2011.",
        "Siciliano B, Khatib O. Springer Handbook of Robotics[M]. Berlin: Springer, 2008.",
    ]
    for idx, ref in enumerate(refs, start=1):
        add_para(doc, f"[{idx}] {ref}")

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
